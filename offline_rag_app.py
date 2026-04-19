# offline_rag_app.py — Zaktualizowana wersja 2025/2026
# Zmiany: nowe importy LangChain, persystencja FAISS, ładowanie JSONL,
#         invoke() zamiast run(), wybór profilu sprzętowego (--profil)

import os
import json
import argparse
import fitz        # PyMuPDF
import docx
import gradio as gr

# === NOWE importy LangChain (community split od v1.x+) ===
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import LlamaCpp
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_classic.chains import RetrievalQA
from api_connectors import SAOSConnector, SejmELIConnector

# ============================================================
#  PROFILE SPRZĘTOWE — wybierasz przy starcie
# ============================================================
PROFILES = {
    "slaby": {
        "opis":        "💻 Słaby komputer (8 GB RAM) — Llama 3.2 8B",
        "model_path":  "Llama-3.2-8B-Instruct-Q4_K_M.gguf",
        "model_url":   "https://huggingface.co/bartowski/Llama-3.2-8B-Instruct-GGUF/resolve/main/Llama-3.2-8B-Instruct-Q4_K_M.gguf",
        "n_ctx":       4096,
        "n_batch":     256,
        "max_tokens":  512,
        "max_orzeczenia": 3_000,
        "max_qa":         1_000,
    },
    "sredni": {
        "opis":        "🖥️ Średni komputer (12–16 GB RAM) — Gemma 4 E4B",
        "model_path":  "gemma-4-e4b-it-Q4_K_M.gguf",
        "model_url":   "https://huggingface.co/bartowski/gemma-4-e4b-it-GGUF/resolve/main/gemma-4-e4b-it-Q4_K_M.gguf",
        "n_ctx":       8192,
        "n_batch":     512,
        "max_tokens":  1024,
        "max_orzeczenia": 7_000,
        "max_qa":         3_000,
    },
    "mocny": {
        "opis":        "🚀 Mocny komputer (16+ GB RAM) — Gemma 4 26B MoE",
        "model_path":  "gemma-4-26b-a4b-it-Q4_K_M.gguf",
        "model_url":   "https://huggingface.co/bartowski/gemma-4-26b-a4b-it-GGUF/resolve/main/gemma-4-26b-a4b-it-Q4_K_M.gguf",
        "n_ctx":       16384,
        "n_batch":     512,
        "max_tokens":  2048,
        "max_orzeczenia": 10_000,
        "max_qa":         5_000,
    },
}

# ==================== USTAWIENIA STAŁE ====================
EMBEDDING_MODEL  = "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
DOCS_FOLDER      = "docs"
JSONL_ORZECZENIA = "orzeczenia_SN.jsonl"
JSONL_QA         = "fine_tune_qa.jsonl"
FAISS_INDEX_PATH = "faiss_index"
# ==========================================================


def wybierz_profil() -> dict:
    """Parsuje --profil z CLI lub pyta interaktywnie."""
    parser = argparse.ArgumentParser(description="Lokalny Asystent Prawniczy RAG")
    parser.add_argument(
        "--profil",
        choices=["slaby", "sredni", "mocny"],
        help="Profil sprzętowy: slaby | sredni | mocny"
    )
    args, _ = parser.parse_known_args()

    if args.profil:
        return PROFILES[args.profil]

    # Interaktywny wybór jeśli nie podano argumentu
    print("\n" + "="*60)
    print("  Lokalny Asystent RAG — Orzeczenia Sądu Najwyższego")
    print("="*60)
    print("\n📋 Wybierz profil sprzętowy:\n")
    opcje = list(PROFILES.keys())
    for i, klucz in enumerate(opcje, 1):
        p = PROFILES[klucz]
        print(f"  [{i}] {p['opis']}")
        print(f"       Model: {p['model_path']}\n")

    while True:
        try:
            wybor = input("Wpisz numer (1/2/3): ").strip()
            if wybor in ["1", "2", "3"]:
                klucz = opcje[int(wybor) - 1]
                print(f"\n✅ Wybrany profil: {PROFILES[klucz]['opis']}\n")
                return PROFILES[klucz]
        except (ValueError, KeyboardInterrupt):
            pass
        print("  Proszę wpisać 1, 2 lub 3.")


# --- Loader: pliki PDF / DOCX / TXT z folderu docs/ ---
def load_folder_documents(folder_path: str) -> list[Document]:
    documents = []
    if not os.path.isdir(folder_path):
        print(f"  [pominięto] folder '{folder_path}' nie istnieje")
        return documents

    for filename in os.listdir(folder_path):
        path = os.path.join(folder_path, filename)
        try:
            if filename.endswith(".pdf"):
                with fitz.open(path) as pdf:
                    text = "\n".join(page.get_text() for page in pdf)
            elif filename.endswith(".docx"):
                doc = docx.Document(path)
                text = "\n".join(p.text for p in doc.paragraphs)
            elif filename.endswith(".txt"):
                with open(path, encoding="utf-8") as f:
                    text = f.read()
            else:
                continue
            documents.append(Document(page_content=text, metadata={"source": filename, "type": "doc"}))
        except Exception as e:
            print(f"  [błąd] {filename}: {e}")

    print(f"  Załadowano {len(documents)} plików z '{folder_path}'")
    return documents


# --- Loader: orzeczenia_SN.jsonl ---
def load_orzeczenia(jsonl_path: str, limit: int | None = None) -> list[Document]:
    documents = []
    if not os.path.isfile(jsonl_path):
        print(f"  [pominięto] plik '{jsonl_path}' nie istnieje")
        return documents

    with open(jsonl_path, encoding="utf-8") as f:
        for i, line in enumerate(f):
            if limit and i >= limit:
                break
            try:
                record = json.loads(line)
                content  = record.get("content", "").strip()
                filename = record.get("filename", f"rekord_{i}")
                if content:
                    documents.append(Document(
                        page_content=content,
                        metadata={"source": filename, "type": "orzeczenie_SN"}
                    ))
            except json.JSONDecodeError:
                continue

    print(f"  Załadowano {len(documents)} orzeczeń z '{jsonl_path}'")
    return documents


# --- Loader: fine_tune_qa.jsonl ---
def load_qa_pairs(jsonl_path: str, limit: int | None = None) -> list[Document]:
    documents = []
    if not os.path.isfile(jsonl_path):
        print(f"  [pominięto] plik '{jsonl_path}' nie istnieje")
        return documents

    with open(jsonl_path, encoding="utf-8") as f:
        for i, line in enumerate(f):
            if limit and i >= limit:
                break
            try:
                record = json.loads(line)
                instruction = record.get("instruction", "").strip()
                context     = record.get("input", "").strip()
                answer      = record.get("output", "").strip()
                text = f"Pytanie: {instruction}\n\nKontekst:\n{context}\n\nOdpowiedź: {answer}"
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": f"qa_{i}", "type": "qa_pair"}
                    ))
            except json.JSONDecodeError:
                continue

    print(f"  Załadowano {len(documents)} par QA z '{jsonl_path}'")
    return documents


# ==================== MAIN ====================

profil = wybierz_profil()
MODEL_PATH = profil["model_path"]

# Unikalny folder indeksu per profil (żeby nie nadpisywać)
profil_nazwa = next(k for k, v in PROFILES.items() if v is profil)
faiss_path   = f"{FAISS_INDEX_PATH}_{profil_nazwa}"

print(f"\n[1] Ładowanie dokumentów (limit: {profil['max_orzeczenia']} orzeczeń)...")
all_docs = []
all_docs += load_folder_documents(DOCS_FOLDER)
all_docs += load_orzeczenia(JSONL_ORZECZENIA,  limit=profil["max_orzeczenia"])
all_docs += load_qa_pairs(JSONL_QA,            limit=profil["max_qa"])
print(f"  Łącznie dokumentów: {len(all_docs)}")

print("\n[2] Tworzenie chunków...")
splitter = RecursiveCharacterTextSplitter(
    chunk_size=600,
    chunk_overlap=80,
    separators=["\n\n", "\n", ". ", " ", ""]
)
chunks = splitter.split_documents(all_docs)
print(f"  Łącznie chunków: {len(chunks)}")

print("\n[3] Embeddingi i FAISS...")
embedding = HuggingFaceEmbeddings(
    model_name=EMBEDDING_MODEL,
    model_kwargs={"device": "cpu"}   # zmień na "cuda" jeśli masz GPU NVIDIA
)

if os.path.isdir(faiss_path):
    print(f"  Wczytywanie istniejącego indeksu z '{faiss_path}'...")
    vectordb = FAISS.load_local(
        faiss_path,
        embedding,
        allow_dangerous_deserialization=True
    )
else:
    print("  Budowanie nowego indeksu (może potrwać kilka minut)...")
    vectordb = FAISS.from_documents(chunks, embedding)
    vectordb.save_local(faiss_path)
    print(f"  Indeks zapisany w '{faiss_path}'")

retriever = vectordb.as_retriever(search_kwargs={"k": 5})

print(f"\n[4] Ładowanie modelu: {MODEL_PATH}")
if not os.path.isfile(MODEL_PATH):
    raise FileNotFoundError(
        f"\n❌ Model '{MODEL_PATH}' nie znaleziony!\n"
        f"Pobierz go z:\n  {profil['model_url']}\n"
        f"i umieść w folderze projektu."
    )

llm = LlamaCpp(
    model_path=MODEL_PATH,
    temperature=0.3,
    max_tokens=profil["max_tokens"],
    top_p=0.95,
    n_ctx=profil["n_ctx"],
    n_batch=profil["n_batch"],
    verbose=False
)

qa_chain = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type="stuff",
    retriever=retriever,
    return_source_documents=True
)

print("\n✅ System gotowy!\n")


# ==================== TRYBY WYSZUKIWANIA ====================

TRYB_LOKALNY  = "🔒 Lokalny  (FAISS — offline, szybki)"
TRYB_ONLINE   = "🌐 Online   (SAOS + Sejm — aktualne dane)"
TRYB_LACZONY  = "🔄 Łącznie  (lokalny + SAOS + Sejm) ⭐"

saos_connector = SAOSConnector()
sejm_connector = SejmELIConnector()


def formatuj_zrodla(docs: list) -> str:
    """Formatuje listę źródeł do wyświetlenia w UI."""
    seen    = set()
    sources = ""
    for doc in docs:
        src  = doc.metadata.get("source", "nieznane")
        typ  = doc.metadata.get("type", "")
        url  = doc.metadata.get("url", "")
        key  = f"{src}|{typ}"
        if key not in seen:
            seen.add(key)
            ikona = {"saos_live": "🌐", "sejm_eli": "📜"}.get(typ, "💾")
            line  = f"{ikona} [{typ}] {src}"
            if url:
                line += f"\n   → {url}"
            sources += line + "\n"
    return sources.strip() or "Brak informacji o źródłach"


def ask_question(query: str, tryb: str) -> tuple[str, str]:
    if not query.strip():
        return "Proszę wpisać pytanie.", ""

    kontekst_docs = []
    status_log    = []

    # --- Ścieżka LOKALNA (FAISS) ---
    if tryb in (TRYB_LOKALNY, TRYB_LACZONY):
        result     = qa_chain.invoke({"query": query})
        faiss_docs = result.get("source_documents", [])
        kontekst_docs.extend(faiss_docs)
        status_log.append(f"💾 Lokalny FAISS: {len(faiss_docs)} chunków")

    # --- Ścieżka ONLINE (SAOS + Sejm ELI) ---
    if tryb in (TRYB_ONLINE, TRYB_LACZONY):
        print(f"\n  🌐 Szukam w SAOS: '{query[:60]}'...")
        saos_docs = saos_connector.search(query, n=4)
        kontekst_docs.extend(saos_docs)
        status_log.append(f"🌐 SAOS: {len(saos_docs)} orzeczeń online")

        print(f"  📜 Szukam w Sejm ELI: '{query[:60]}'...")
        sejm_docs = sejm_connector.search(query, n=3)
        kontekst_docs.extend(sejm_docs)
        status_log.append(f"📜 Sejm ELI: {len(sejm_docs)} aktów prawnych")

    if not kontekst_docs:
        return "Brak wyników — sprawdź połączenie z internetem lub zmień tryb na Lokalny.", ""

    # --- Dla trybu ONLINE nie używamy RetrievalQA — budujemy prompt ręcznie ---
    if tryb == TRYB_ONLINE:
        kontekst_text = "\n\n---\n\n".join(
            doc.page_content for doc in kontekst_docs
        )
        prompt = (
            f"Na podstawie poniższych dokumentów odpowiedz na pytanie.\n"
            f"Pytanie: {query}\n\n"
            f"Dokumenty:\n{kontekst_text[:profil['n_ctx'] - 500]}\n\n"
            f"Odpowiedź:"
        )
        answer = llm.invoke(prompt)

    elif tryb == TRYB_LACZONY:
        # Łączony: FAISS dał już odpowiedź + dorzucamy kontekst z API
        faiss_result = qa_chain.invoke({"query": query})
        faiss_answer = faiss_result.get("result", "")

        # Budujemy dodatkowy kontekst tylko z API
        api_docs    = [d for d in kontekst_docs if d.metadata.get("type") in ("saos_live", "sejm_eli")]
        api_context = "\n\n".join(d.page_content for d in api_docs)[:1500]

        if api_docs:
            prompt = (
                f"Odpowiedź z lokalnej bazy orzeczeń:\n{faiss_answer}\n\n"
                f"Uzupełnij odpowiedź o poniższe aktualne dane z SAOS/Sejm:\n{api_context}\n\n"
                f"Pytanie: {query}\nZintegrowana odpowiedź:"
            )
            answer = llm.invoke(prompt)
        else:
            answer = faiss_answer

    else:  # TRYB_LOKALNY
        result = qa_chain.invoke({"query": query})
        answer = result.get("result", "Brak odpowiedzi")

    sources = formatuj_zrodla(kontekst_docs)
    # Dopisujemy logi statusu na dole źródeł
    sources += "\n\n" + " | ".join(status_log)

    return answer, sources


# ==================== INTERFEJS GRADIO ====================

with gr.Blocks(
    theme=gr.themes.Soft(),
    title="RAG — Orzeczenia Sądu Najwyższego",
    css="""
        .tryb-radio label { font-size: 0.95em !important; }
        footer { display: none !important; }
    """
) as interface:

    gr.Markdown(f"""
    # ⚖️ Lokalny Asystent Prawniczy RAG
    **Profil:** {profil['opis']} &nbsp;|&nbsp; Baza: Orzeczenia SN + Dziennik Ustaw
    """)

    with gr.Row():
        tryb_radio = gr.Radio(
            choices=[TRYB_LOKALNY, TRYB_ONLINE, TRYB_LACZONY],
            value=TRYB_LOKALNY,
            label="📡 Tryb wyszukiwania",
            elem_classes="tryb-radio",
        )

    gr.Markdown("""
    > **🔒 Lokalny** — szybki, działa bez netu, korzysta z pobranych orzeczeń SN  
    > **🌐 Online** — pobiera aktualne orzeczenia z SAOS i przepisy z Dziennika Ustaw (~3-5 sek więcej)  
    > **🔄 Łącznie** — łączy oba podejścia, najlepsza odpowiedź ⭐
    """)

    with gr.Row():
        with gr.Column(scale=3):
            query_box = gr.Textbox(
                label="Twoje pytanie",
                placeholder="np. Kiedy Sąd Najwyższy odmawia przyjęcia skargi kasacyjnej?",
                lines=3
            )
            submit_btn = gr.Button("🔍 Szukaj", variant="primary", size="lg")

        with gr.Column(scale=2):
            source_box = gr.Textbox(
                label="📂 Źródła",
                lines=8,
                interactive=False
            )

    answer_box = gr.Textbox(
        label="📋 Odpowiedź",
        lines=12,
        interactive=False
    )

    submit_btn.click(
        fn=ask_question,
        inputs=[query_box, tryb_radio],
        outputs=[answer_box, source_box]
    )

    gr.Examples(
        examples=[
            ["Kiedy Sąd Najwyższy odmawia przyjęcia skargi kasacyjnej do rozpoznania?"],
            ["Co to znaczy orzeczenie niezgodne z prawem w rozumieniu art. 4241 kpc?"],
            ["Jakie są przesłanki podziału majątku wspólnego małżonków?"],
            ["Znajdź aktualny tekst Kodeksu postępowania cywilnego"],
            ["Jakie przepisy regulują skargę kasacyjną?"],
        ],
        inputs=query_box
    )

interface.launch(server_name="0.0.0.0", server_port=7860)
